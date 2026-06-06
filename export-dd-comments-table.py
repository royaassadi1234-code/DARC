from __future__ import annotations

import re
from difflib import SequenceMatcher
from pathlib import Path

import fitz
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
PDF_PATH = ROOT / "dd 1.pdf"
OUT_PATH = ROOT / "dd-comments-aligned.docx"
DD_EN_PATH = ROOT / "DD-en.txt"
DD_SOURCE_PATH = ROOT / "Dd.txt"


def annotation_words(page: fitz.Page, annot: fitz.Annot) -> str:
    words = page.get_text("words")
    quad_rects = get_quad_rects(annot)
    if not quad_rects:
        quad_rects = [annot.rect]

    selected = []
    for word in words:
        x0, y0, x1, y1, text, block, line, number = word[:8]
        center = fitz.Point((x0 + x1) / 2, (y0 + y1) / 2)
        if any(rect.contains(center) for rect in quad_rects):
            selected.append((block, line, number, x0, y0, text))

    selected.sort(key=lambda item: (item[0], item[1], item[2], item[3]))
    return clean_text(" ".join(item[5] for item in selected))


def get_quad_rects(annot: fitz.Annot) -> list[fitz.Rect]:
    vertices = getattr(annot, "vertices", None) or []
    rects = []
    for index in range(0, len(vertices), 4):
        points = vertices[index : index + 4]
        if len(points) != 4:
            continue
        xs = [point[0] if isinstance(point, tuple) else point.x for point in points]
        ys = [point[1] if isinstance(point, tuple) else point.y for point in points]
        rect = fitz.Rect(min(xs), min(ys), max(xs), max(ys))
        rect.x0 -= 1.5
        rect.x1 += 1.5
        rect.y0 -= 1.5
        rect.y1 += 1.5
        rects.append(rect)
    return rects


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def section_from_text(value: str) -> str:
    value = value.replace("\\.", ".")
    match = re.search(r"\((?:Int\.?\s*|In[tU]\.?\s*)?([0-9OILS]+|[IVX]+)\s*[\.:,\\]?\s*([0-9OILS]+)?\)", value, re.I)
    if not match:
        regular = re.search(r"\(([0-9OILS]+)\s*[\.:,\\]\s*([0-9OILS]+)\)", value, re.I)
        if regular:
            return f"{ocr_number(regular.group(1))}.{ocr_number(regular.group(2))}"
        return ""

    first, second = match.groups()
    if value[match.start() : match.end()].lower().startswith("(int"):
        return f"0.{ocr_number(first)}"
    if second:
        if ocr_number(first) == "3" and ocr_number(second) == "10" and "renovation" in value.lower():
            return "31.10"
        return f"{ocr_number(first)}.{ocr_number(second)}"
    return ""


def section_near_annotation(page: fitz.Page, annot: fitz.Annot) -> str:
    right_column = fitz.Rect(page.rect.width * 0.50, 0, page.rect.width, annot.rect.y1 + 4)
    text_before = clean_text(page.get_textbox(right_column))
    markers = re.findall(r"\((?:Int\.?\s*|In[tU]\.?\s*)?[0-9OILSIVX]+\s*[\.:,\\]?\s*[0-9OILS]*\)", text_before, re.I)
    for marker in reversed(markers):
        section = section_from_text(marker)
        if section:
            return section
    return ""


def ocr_number(value: str) -> str:
    return (
        value.upper()
        .replace("O", "0")
        .replace("I", "1")
        .replace("L", "1")
        .replace("S", "5")
    )


def extract_annotations() -> list[dict[str, object]]:
    doc = fitz.open(PDF_PATH)
    dd_translations = load_dd_translations()
    dd_source = load_dd_source_records()
    comments = []
    left_highlights_by_page: dict[int, list[dict[str, object]]] = {}

    for page_index, page in enumerate(doc):
        annot = page.first_annot
        while annot:
            if annot.type[1] != "Highlight":
                annot = annot.next
                continue

            info = annot.info or {}
            content = clean_text(info.get("content") or "")
            text = annotation_words(page, annot) or clean_text(page.get_textbox(annot.rect))
            item = {
                "page": page_index + 1,
                "rect": annot.rect,
                "text": text,
                "comment": content,
                "section": section_from_text(text) or section_near_annotation(page, annot),
            }
            if not item["section"]:
                item["section"] = section_from_translation_text(text, dd_translations)

            if content and annot.rect.x0 > page.rect.width * 0.52:
                comments.append(item)
            elif annot.rect.x1 < page.rect.width * 0.52:
                left_highlights_by_page.setdefault(page_index + 1, []).append(item)

            annot = annot.next

    for comment in comments:
        raw_transliteration = align_left_highlights(comment, left_highlights_by_page.get(comment["page"], []))
        comment["transliteration"] = correct_transliteration_highlights(
            raw_transliteration,
            str(comment.get("section") or ""),
            dd_source,
        )

    return comments


def load_dd_source_records() -> dict[str, str]:
    records: dict[str, str] = {}
    current_location = ""
    for line in DD_SOURCE_PATH.read_text(encoding="utf-8").splitlines():
        trimmed = line.strip()
        if not trimmed:
            continue
        section = re.match(r"^([0-9]+(?:\.[0-9]+[a-z]?)?)\s+(.*)$", trimmed)
        if section:
            current_location = section.group(1)
            records[current_location] = section.group(2)
            continue
        if current_location:
            records[current_location] = f"{records[current_location]} {trimmed}"
    return records


def load_dd_translations() -> list[tuple[str, str]]:
    if not DD_EN_PATH.exists():
        return []
    records = []
    for line in DD_EN_PATH.read_text(encoding="utf-8").splitlines():
        if "\t" not in line:
            continue
        location, text = line.split("\t", 1)
        records.append((location, text))
    return records


def section_from_translation_text(value: str, records: list[tuple[str, str]]) -> str:
    query = normalize_for_match(value)
    if len(query) < 30:
        return ""

    best_score = 0.0
    best_location = ""
    for location, text in records:
        target = normalize_for_match(text)
        if not target:
            continue
        score = SequenceMatcher(None, query[:420], target[:900]).ratio()
        if query[:100] and query[:100] in target:
            score += 1
        if score > best_score:
            best_score = score
            best_location = location
    return best_location if best_score >= 0.72 else ""


def normalize_for_match(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def align_left_highlights(comment: dict[str, object], left_items: list[dict[str, object]]) -> str:
    if not left_items:
        return ""
    rect: fitz.Rect = comment["rect"]  # type: ignore[assignment]
    center = (rect.y0 + rect.y1) / 2
    height = max(rect.height, 36)
    tolerance = max(56, height * 1.35)

    nearby = [
        item for item in left_items
        if abs((((item["rect"].y0 + item["rect"].y1) / 2) - center)) <= tolerance  # type: ignore[index, union-attr]
    ]
    if not nearby:
        nearby = sorted(left_items, key=lambda item: abs((((item["rect"].y0 + item["rect"].y1) / 2) - center)))[:6]  # type: ignore[index, union-attr]

    nearby.sort(key=lambda item: (item["rect"].y0, item["rect"].x0))  # type: ignore[index, union-attr]
    words = []
    for item in nearby:
        text = str(item["text"])
        if text and text not in words:
            words.append(text)
    return "; ".join(words)


def correct_transliteration_highlights(raw: str, section: str, dd_source: dict[str, str]) -> str:
    if not raw or not section:
        return raw
    source_text = source_text_for_section(section, dd_source)
    if not source_text:
        return raw

    corrected = []
    for part in [item.strip() for item in raw.split(";") if item.strip()]:
        corrected_part = closest_source_phrase(part, source_text)
        if corrected_part and corrected_part not in corrected:
            corrected.append(corrected_part)
        elif part not in corrected:
            corrected.append(part)
    return "; ".join(corrected)


def source_text_for_section(section: str, dd_source: dict[str, str]) -> str:
    nearby = nearby_source_text(section, dd_source)
    if nearby:
        return nearby
    if section in dd_source:
        return dd_source[section]
    related = [
        text for location, text in dd_source.items()
        if location.startswith(f"{section}") and re.match(rf"^{re.escape(section)}[a-z]$", location)
    ]
    if related:
        return " ".join(related)
    chapter, _, stanza = section.partition(".")
    if chapter and stanza:
        related = [
            text for location, text in dd_source.items()
            if location.startswith(f"{chapter}.{stanza}") and re.match(rf"^{re.escape(chapter)}\.{re.escape(stanza)}[a-z]$", location)
        ]
        if related:
            return " ".join(related)
    return ""


def nearby_source_text(section: str, dd_source: dict[str, str]) -> str:
    match = re.match(r"^(\d+)\.(\d+)", section)
    if not match:
        return ""
    chapter = int(match.group(1))
    stanza = int(match.group(2))
    wanted = {
        f"{chapter}.{number}"
        for number in range(max(0, stanza - 2), stanza + 4)
    }
    chunks = []
    for location, text in dd_source.items():
        base = re.sub(r"[a-z]$", "", location)
        if base in wanted:
            chunks.append(text)
    return " ".join(chunks)


def closest_source_phrase(ocr_phrase: str, source_text: str) -> str:
    source_tokens = tokenize_transliteration(source_text)
    query_tokens = tokenize_transliteration(ocr_phrase)
    if not source_tokens or not query_tokens:
        return ""

    query_folded = fold_transliteration(" ".join(query_tokens))
    best_score = 0.0
    best_phrase = ""
    min_size = max(1, len(query_tokens) - 1)
    max_size = min(len(source_tokens), len(query_tokens) + 2)

    for size in range(min_size, max_size + 1):
        for start in range(0, len(source_tokens) - size + 1):
            phrase_tokens = source_tokens[start : start + size]
            phrase = " ".join(phrase_tokens)
            phrase_folded = fold_transliteration(phrase)
            score = SequenceMatcher(None, query_folded, phrase_folded).ratio()
            if query_folded and (query_folded in phrase_folded or phrase_folded in query_folded):
                score += 0.35
            if score > best_score:
                best_score = score
                best_phrase = phrase

    if best_score >= 0.58:
        return best_phrase

    corrected_tokens = []
    for token in query_tokens:
        corrected = closest_source_token(token, source_tokens)
        if corrected:
            corrected_tokens.append(corrected)
    return " ".join(corrected_tokens) if corrected_tokens else ocr_phrase


def closest_source_token(ocr_token: str, source_tokens: list[str]) -> str:
    query = fold_transliteration(ocr_token)
    if not query:
        return ""
    best_score = 0.0
    best_token = ""
    for token in source_tokens:
        folded = fold_transliteration(token)
        score = SequenceMatcher(None, query, folded).ratio()
        if query in folded or folded in query:
            score += 0.25
        if score > best_score:
            best_score = score
            best_token = token
    return best_token if best_score >= 0.62 else ocr_token


def tokenize_transliteration(value: str) -> list[str]:
    cleaned = re.sub(r"\[[^\]]+\]|<[^>]+>", " ", value)
    return re.findall(r"[\wāēīōūčšǰγθδĀĒĪŌŪČŠǰ-]+", cleaned, flags=re.UNICODE)


def fold_transliteration(value: str) -> str:
    replacements = {
        "ā": "a",
        "ē": "e",
        "ī": "i",
        "ō": "o",
        "ū": "u",
        "č": "c",
        "š": "s",
        "ǰ": "j",
        "γ": "g",
        "θ": "t",
        "δ": "d",
    }
    text = value.lower()
    for source, target in replacements.items():
        text = text.replace(source, target)
    text = text.replace("ı", "i")
    return re.sub(r"[^a-z0-9-]+", " ", text).strip()


def build_docx(rows: list[dict[str, object]]) -> None:
    document = Document()
    document.core_properties.title = "DD annotated translation comments aligned with transliteration highlights"
    document.add_heading("DD annotated translation comments", level=1)
    intro = document.add_paragraph(
        "Extracted from dd 1.pdf. Rows contain comments on highlighted translation passages and nearby highlighted words in the transliteration column."
    )
    intro.style.font.size = Pt(10)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Chapter.Stanza", "PDF page", "Highlighted transliteration word(s)", "Highlighted translation passage", "Comment"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True

    for row in rows:
        cells = table.add_row().cells
        values = [
            str(row.get("section") or ""),
            str(row.get("page") or ""),
            str(row.get("transliteration") or ""),
            str(row.get("text") or ""),
            str(row.get("comment") or ""),
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if contains_rtl(value) else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    document.save(OUT_PATH)


def contains_rtl(value: str) -> bool:
    return bool(re.search(r"[\u0600-\u06ff]", value))


def main() -> None:
    rows = extract_annotations()
    build_docx(rows)
    print(f"Exported {len(rows)} rows to {OUT_PATH.name}")


if __name__ == "__main__":
    main()
