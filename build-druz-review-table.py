from __future__ import annotations

import json
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DD_TEXT = Path("Dd.txt")
COMMENTS_JSON = Path("dd-comment-highlights.json")
PDF_TRANSLATIONS_JSON = Path("dd-pdf-translations.json")
PERSIAN_TRANSLATIONS_JSON = Path("dd2mahshid-persian-translations.json")
OUTPUT = Path("pdf-output/druz-review-table.docx")
FALLBACK_OUTPUT = Path("pdf-output/druz-review-table-updated.docx")

DRUZ_VARIANTS = [
    "druz",
    "druj",
    "drug",
    "drux",
    "druzān",
    "druzan",
    "druzīh",
    "druzih",
    "druxtārīh",
    "druxtarih",
    "druzaskān",
    "druzaskan",
]


def fold_text(value: str) -> str:
    replacements = str.maketrans({
        "ā": "a",
        "ē": "e",
        "ī": "i",
        "ō": "o",
        "ū": "u",
        "š": "s",
        "č": "c",
        "ǰ": "j",
        "γ": "g",
        "θ": "t",
        "Ā": "a",
        "Ē": "e",
        "Ī": "i",
        "Ō": "o",
        "Ū": "u",
        "Š": "s",
        "Č": "c",
        "Ǧ": "g",
    })
    return value.translate(replacements).lower()


def parse_records(raw: str) -> list[dict[str, str]]:
    records = []
    current = None
    section_pattern = re.compile(r"^([0-9]+(?:\.[0-9]+[a-z]?)?)\s+(.+)$")

    for index, line in enumerate(raw.splitlines()):
        trimmed = line.strip()
        if not trimmed:
            continue

        section = section_pattern.match(trimmed)
        if section:
            current = {
                "index": index,
                "location": section.group(1),
                "text": section.group(2),
            }
            records.append(current)
            continue

        if current:
            current["text"] += f" {trimmed}"
        else:
            current = {
                "index": index,
                "location": f"line {index + 1}",
                "text": trimmed,
            }
            records.append(current)

    return records


def base_location(location: str) -> str:
    return re.sub(r"[a-z]+$", "", str(location).strip())


def load_comments() -> dict[str, list[dict[str, object]]]:
    comments = json.loads(COMMENTS_JSON.read_text(encoding="utf-8"))
    by_location = defaultdict(list)
    for comment in comments:
        by_location[str(comment.get("location", "")).strip()].append(comment)
    return by_location


def load_pdf_translations() -> dict[str, dict[str, object]]:
    if not PDF_TRANSLATIONS_JSON.exists():
        return {}
    return json.loads(PDF_TRANSLATIONS_JSON.read_text(encoding="utf-8"))


def load_persian_translations() -> dict[str, dict[str, object]]:
    if not PERSIAN_TRANSLATIONS_JSON.exists():
        return {}
    return json.loads(PERSIAN_TRANSLATIONS_JSON.read_text(encoding="utf-8"))


def variant_ranges(text: str) -> list[tuple[int, int]]:
    folded = fold_text(text)
    ranges = []
    boundary = r"[\w=-]"

    for variant in DRUZ_VARIANTS:
        pattern = re.compile(rf"(?<!{boundary}){re.escape(fold_text(variant))}(?!{boundary})")
        for match in pattern.finditer(folded):
            ranges.append((match.start(), match.end()))

    return merge_ranges(ranges)


def merge_ranges(ranges: list[tuple[int, int]]) -> list[tuple[int, int]]:
    merged = []
    for start, end in sorted(ranges):
        if merged and start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    return merged


def matched_words(text: str) -> list[str]:
    return sorted({text[start:end] for start, end in variant_ranges(text)}, key=str.lower)


def add_highlighted_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    ranges = variant_ranges(text)
    if not ranges:
        paragraph.add_run(text)
        return

    cursor = 0
    for start, end in ranges:
        if start > cursor:
            paragraph.add_run(text[cursor:start])
        run = paragraph.add_run(text[start:end])
        run.bold = True
        cursor = end
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_dd_paragraph_with_translation(
    cell,
    paragraph_text: str,
    translation_passage: str,
    translation_source: str,
    persian_translation: str,
    persian_source: str,
) -> None:
    add_highlighted_text(cell, paragraph_text)
    if translation_passage:
        label_paragraph = cell.add_paragraph()
        label = "Highlighted translation passage:" if translation_source == "comments" else "Translation from dd 1.pdf:"
        label_run = label_paragraph.add_run(label)
        label_run.bold = True
        label_run.italic = True

        translation_paragraph = cell.add_paragraph()
        translation_run = translation_paragraph.add_run(translation_passage)
        translation_run.italic = True

    if not persian_translation:
        return

    persian_label_paragraph = cell.add_paragraph()
    persian_label_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    persian_label_run = persian_label_paragraph.add_run(persian_source)
    persian_label_run.bold = True
    persian_label_run.italic = True

    persian_paragraph = cell.add_paragraph()
    persian_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    persian_run = persian_paragraph.add_run(persian_translation)
    persian_run.italic = True


def combined_comment_field(comments: list[dict[str, object]], key: str) -> str:
    values = []
    for comment in comments:
        value = comment.get(key, "")
        if isinstance(value, list):
            value = "; ".join(str(item) for item in value if item)
        value = str(value).strip()
        if value:
            values.append(value)
    return "\n\n".join(values)


def matching_comments(record: dict[str, str], comments_by_location: dict[str, list[dict[str, object]]]) -> list[dict[str, object]]:
    location = record["location"]
    comments = []
    comments.extend(comments_by_location.get(location, []))
    if location != base_location(location):
        comments.extend(comments_by_location.get(base_location(location), []))

    seen = set()
    unique = []
    for comment in comments:
        identity = (
            comment.get("location"),
            comment.get("title"),
            comment.get("englishComment"),
        )
        if identity not in seen:
            seen.add(identity)
            unique.append(comment)
    return unique


def matching_pdf_translation(
    record: dict[str, str],
    pdf_translations: dict[str, dict[str, object]],
) -> dict[str, object] | None:
    location = record["location"]
    return pdf_translations.get(location) or pdf_translations.get(base_location(location))


def matching_persian_translation(
    record: dict[str, str],
    persian_translations: dict[str, dict[str, object]],
) -> dict[str, object] | None:
    chapter = base_location(record["location"]).split(".", 1)[0]
    return persian_translations.get(chapter)


def set_cell_text(cell, text: str, rtl: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if rtl:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run(text)


def build_document(rows: list[dict[str, object]]) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    document.add_heading("Druz Review Table", level=1)
    intro = document.add_paragraph()
    intro.add_run("Source: ").bold = True
    intro.add_run("Dd.txt paragraphs containing Druz variants, merged with dd-comments-aligned-with-english.docx data.")
    document.add_paragraph(f"Rows: {len(rows)}")

    headers = [
        "DD location",
        "Matched Druz form(s)",
        "DD paragraph + highlighted translation passage",
        "PDF page",
        "Highlighted/corrected words",
        "Comment title",
        "English comment",
        "Persian comment",
        "Review note",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    widths = [0.7, 1.0, 4.8, 0.6, 1.7, 1.6, 3.0, 2.5, 1.4]

    for row_data in rows:
        cells = table.add_row().cells
        for index, width in enumerate(widths):
            cells[index].width = Inches(width)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        set_cell_text(cells[0], str(row_data["location"]))
        set_cell_text(cells[1], "; ".join(row_data["matchedWords"]))
        add_dd_paragraph_with_translation(
            cells[2],
            str(row_data["paragraph"]),
            str(row_data["translationPassage"]),
            str(row_data["translationSource"]),
            str(row_data["persianTranslation"]),
            str(row_data["persianTranslationSource"]),
        )
        set_cell_text(cells[3], str(row_data["pdfPage"]))
        set_cell_text(cells[4], str(row_data["highlightedWords"]))
        set_cell_text(cells[5], str(row_data["commentTitle"]))
        set_cell_text(cells[6], str(row_data["englishComment"]))
        set_cell_text(cells[7], str(row_data["persianComment"]), rtl=True)
        set_cell_text(cells[8], "")

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        document.save(OUTPUT)
        return OUTPUT
    except PermissionError:
        document.save(FALLBACK_OUTPUT)
        return FALLBACK_OUTPUT


def main() -> None:
    records = parse_records(DD_TEXT.read_text(encoding="utf-8"))
    comments_by_location = load_comments()
    pdf_translations = load_pdf_translations()
    persian_translations = load_persian_translations()
    rows = []

    for record in records:
        words = matched_words(record["text"])
        if not words:
            continue

        comments = matching_comments(record, comments_by_location)
        comment_translation = combined_comment_field(comments, "translationPassage")
        pdf_translation = matching_pdf_translation(record, pdf_translations)
        translation_passage = comment_translation
        translation_source = "comments" if comment_translation else ""
        pdf_pages = combined_comment_field(comments, "pdfPage")

        if not translation_passage and pdf_translation:
            translation_passage = str(pdf_translation.get("translation", ""))
            translation_source = "pdf"
            pages = pdf_translation.get("pdfPages", [])
            if pages:
                pdf_pages = "; ".join(str(page) for page in pages)

        persian_translation = matching_persian_translation(record, persian_translations)
        persian_translation_text = ""
        persian_translation_source = ""
        if persian_translation:
            persian_translation_text = str(persian_translation.get("text", ""))
            pages = persian_translation.get("pdfPages", [])
            page_label = ", ".join(str(page) for page in pages)
            persian_translation_source = f"Persian translation from DD2Mahshid.PDF OCR, PDF page(s): {page_label}"

        rows.append({
            "location": record["location"],
            "matchedWords": words,
            "paragraph": record["text"],
            "pdfPage": pdf_pages,
            "highlightedWords": combined_comment_field(comments, "highlightedWords"),
            "commentTitle": combined_comment_field(comments, "title"),
            "englishComment": combined_comment_field(comments, "englishComment"),
            "persianComment": combined_comment_field(comments, "persianComment"),
            "translationPassage": translation_passage,
            "translationSource": translation_source,
            "persianTranslation": persian_translation_text,
            "persianTranslationSource": persian_translation_source,
        })

    output = build_document(rows)
    print(f"Wrote {output} with {len(rows)} rows.")


if __name__ == "__main__":
    main()
