from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SOURCE = Path("Dd.txt")
DOCX_TARGET = Path("dd-ch41-end-translation-worksheet.docx")
JSON_TARGET = Path("dd-ch41-end-translation-template.json")

PASSAGE_RE = re.compile(r"(?m)^([0-9]+\.[0-9]+[a-z]?)\t")
LOCATION_RE = re.compile(r"\[([^\]]+)\]")


def iter_passages(text: str) -> list[dict[str, object]]:
    matches = list(PASSAGE_RE.finditer(text))
    passages: list[dict[str, object]] = []

    for index, match in enumerate(matches):
        passage_id = match.group(1)
        chapter = int(passage_id.split(".", 1)[0])
        if chapter < 41:
            continue

        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        body = text[match.end() : end].strip()
        locations = sorted(dict.fromkeys(LOCATION_RE.findall(body)))
        passages.append(
            {
                "passageId": passage_id,
                "chapter": chapter,
                "locations": locations,
                "sourceText": body,
                "persianTranslation": "",
                "notes": "",
            }
        )

    return passages


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.xpath("./w:tblHeader")
    if not tbl_header:
        from docx.oxml import OxmlElement

        element = OxmlElement("w:tblHeader")
        element.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "true")
        tr_pr.append(element)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_multiline_text(cell, text: str, *, size: int = 8) -> None:
    cell.text = ""
    lines = text.splitlines() or [""]
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        run.font.size = Pt(size)


def build_docx(passages: list[dict[str, object]]) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)

    title = document.add_heading("Dadestan i Denig Translation Worksheet: Chapters 41-93", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = document.add_paragraph()
    summary.add_run("Purpose: ").bold = True
    summary.add_run(
        "Use the empty Persian translation column to add the corresponding Persian text for each passage. "
        "Keep the Passage ID unchanged so translations can be merged back into the project later."
    )
    summary.paragraph_format.space_after = Pt(8)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Passage ID", "Locations", "Source passage", "Persian translation", "Notes"]
    widths = [0.8, 1.35, 4.2, 4.2, 1.15]

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_text(cell, header, bold=True, size=8)
        cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for passage in passages:
        row = table.add_row()
        row.height = Inches(0.45)
        cells = row.cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        cells[2].width = Inches(widths[2])
        cells[3].width = Inches(widths[3])
        cells[4].width = Inches(widths[4])

        set_cell_text(cells[0], str(passage["passageId"]), bold=True, size=8)
        set_cell_text(cells[1], ", ".join(passage["locations"]) or "-", size=7)
        add_multiline_text(cells[2], str(passage["sourceText"]), size=7)
        set_cell_text(cells[3], "", size=9)
        set_cell_text(cells[4], "", size=8)

        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    document.add_page_break()
    appendix = document.add_heading("Chapter Index", level=2)
    appendix.alignment = WD_ALIGN_PARAGRAPH.LEFT
    counts: dict[int, int] = {}
    for passage in passages:
        counts[int(passage["chapter"])] = counts.get(int(passage["chapter"]), 0) + 1
    for chapter in sorted(counts):
        document.add_paragraph(f"Chapter {chapter}: {counts[chapter]} passages")

    document.save(DOCX_TARGET)


def build_json(passages: list[dict[str, object]]) -> None:
    JSON_TARGET.write_text(
        json.dumps(
            {
                "source": str(SOURCE),
                "range": "41-end",
                "passageCount": len(passages),
                "passages": passages,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    passages = iter_passages(text)
    build_docx(passages)
    build_json(passages)
    print(f"Wrote {DOCX_TARGET} and {JSON_TARGET} with {len(passages)} passages.")


if __name__ == "__main__":
    main()
